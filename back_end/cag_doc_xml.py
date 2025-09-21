
from docx import Document
from lxml import etree
from docx.oxml.ns import qn
import os
import logging
import re
from docx_recovery_tool import validate_and_repair_docx

# Setup logger
logger = logging.getLogger("doc_xml_colspan")
logger.setLevel(logging.INFO)
fh = logging.FileHandler("doc_xml_colspan.log")
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
fh.setFormatter(formatter)
logger.addHandler(fh)

# Enhanced budget section heading patterns
budget_heading_patterns = [
    r"BUDGET\s*/?\s*FINANCIAL\s+PERFORMANCE",
    r"FINANCIAL\s+PERFORMANCE",
    r"BUDGET\s+PERFORMANCE", 
    r"BUDGET\s+AND\s+FINANCIAL\s+PERFORMANCE",
    r"c\)\s*Budget\s+and\s+Financial\s+Performance",
    r"^\s*c\)\s*Budget",
    r"Financial\s+Performance:",
    r"Budget\s+and\s+Financial\s+Performance:",
    r"Budget\s+allocation\s+for\s+the\s+audit\s+period:",
    r"^\s*Financial\s+Performance\s*:",
    r"^\s*Budget\s+and\s+Financial\s+Performance\s*:",
    r"^\s*Budget\s+allocation\s+for\s+the\s+audit\s+period\s*:",
    r"^\s*\d+\.\d+\s+Financial\s+Performance\s*:",
    r"^\s*\d+\.\d+\s*\t+Financial\s+Performance\s*:",
    r"\d+\.\d+\s*Financial\s+Performance:",
    r"\d+\.\d+\s*\t+Financial\s+Performance:"
]

objective_patterns = [
    r"Audit\s+objectives",
    r"AUDIT\s+OBJECTIVE", 
    r"Audit\s+objective",
    r"AUDIT\s+OBJECTIVE:",
    r"Audit\s+objectives:"
]

criteria_patterns = [
    r"Audit\s+criteria",
    r"AUDIT\s+CRITERIA", 
    r"Audit\s+criteria:",
    r"AUDIT\s+CRITERIA:"
]

def matches_heading_patterns(text):
    """
    Check if the text matches any of the budget, objective, or criteria patterns.
    Returns True if a pattern is found, False otherwise.
    """
    if not text or not text.strip():
        return False
    
    text_stripped = text.strip()
    
    # Check all pattern lists
    all_patterns = budget_heading_patterns + objective_patterns + criteria_patterns
    
    for pattern in all_patterns:
        if re.search(pattern, text_stripped, re.IGNORECASE):
            logger.info(f"Pattern matched: '{pattern}' in text: '{text_stripped}'")
            return True
    
    return False

def is_in_part_one_context(text):
    """
    Check if the text indicates we're in Part I context.
    """
    if not text or not text.strip():
        return False
    
    text_stripped = text.strip()
    part_one_patterns = [
        r"PART\s*I",
        r"Part\s*I",
        r"PART\s*1",
        r"Part\s*1"
    ]
    
    for pattern in part_one_patterns:
        if re.search(pattern, text_stripped, re.IGNORECASE):
            return True
    
    return False


def docx_to_custom_xml(docx_path, xml_output_path):
    logger.info(f"Processing DOCX: {docx_path}")
    
    # Validate and repair DOCX file if needed
    try:
        is_valid, working_docx_path = validate_and_repair_docx(docx_path, repair_if_needed=True)
        if not is_valid:
            logger.error(f"DOCX file validation failed and could not be repaired: {docx_path}")
            return
        
        # working_docx_path will be the same as docx_path since we repair in-place
        doc = Document(working_docx_path)
        
    except Exception as e:
        logger.error(f"Failed to open DOCX file {docx_path}: \"{e}\"")
        return
    
    root = etree.Element("document")

    def process_table(parent_el, table):
        logger.info(f"Processing table at parent: {parent_el.tag}")
        """
        Processes a docx table, creating a robust XML structure that handles
        merged cells (colspan & rowspan) and nested tables, preventing duplication.
        """
        table_el = etree.SubElement(parent_el, "table")

        # A grid to mark cells that are part of a span and already processed
        processed_cells = set()

        # Handle empty tables gracefully
        if not table.rows:
            return
        
        # Handle tables with empty rows gracefully
        row_cell_counts = [len(row.cells) for row in table.rows if row.cells]
        if not row_cell_counts:
            return
            
        n_cols = max(row_cell_counts)
        logger.info(f"Table has {len(table.rows)} rows and {n_cols} columns")
        for r, row in enumerate(table.rows):
            row_el = etree.SubElement(table_el, "row")
            logger.debug(f"Processing row {r+1}/{len(table.rows)}")
            c = 0
            while c < n_cols:
                if (r, c) in processed_cells:
                    c += 1
                    continue

                try:
                    cell = row.cells[c]
                except IndexError:
                    logger.warning(f"IndexError: Row {r}, Col {c} out of range.")
                    break


                # Get cell's merge properties
                tcPr = cell._tc.get_or_add_tcPr()

                # Colspan
                gridSpan_el = tcPr.find(qn('w:gridSpan'))
                colspan = int(gridSpan_el.get(qn('w:val'))) if gridSpan_el is not None else 1

                # Rowspan
                vMerge_el = tcPr.find(qn('w:vMerge'))
                rowspan = 1
                if vMerge_el is not None and vMerge_el.get(qn('w:val')) == 'restart':
                    for i in range(r + 1, len(table.rows)):
                        try:
                            next_cell = table.cell(i, c)
                            next_vMerge_el = next_cell._tc.get_or_add_tcPr().find(qn('w:vMerge'))
                            if next_vMerge_el is not None and next_vMerge_el.get(qn('w:val')) is None:
                                rowspan += 1
                            else:
                                break
                        except IndexError:
                            logger.warning(f"IndexError in rowspan calculation at row {i}, col {c}")
                            break


                # Mark all cells covered by this span as processed
                for i in range(r, r + rowspan):
                    for j in range(c, c + colspan):
                        processed_cells.add((i, j))


                # Create the XML element for the cell
                cell_el = etree.SubElement(row_el, "cell")
                if rowspan > 1:
                    cell_el.set('rowspan', str(rowspan))
                if colspan > 1:
                    cell_el.set('colspan', str(colspan))

                # Add content to the cell
                cell_text = "\n".join(p.text for p in cell.paragraphs)
                cell_el.text = cell_text
                logger.debug(f"Cell at row {r}, col {c}: rowspan={rowspan}, colspan={colspan}, text='{cell_text[:30]}...'")
                for nested_table in cell.tables:
                    logger.info(f"Processing nested table in cell at row {r}, col {c}")
                    process_table(cell_el, nested_table)

                c += colspan

    # --- Main document processing ---
    
    # Track if we're in Part I context
    in_part_one = False

    for block in doc.element.body:
        if block.tag.endswith('p'):
            para = next((p for p in doc.paragraphs if p._element == block), None)
            if para is None: continue

            style = para.style.name.lower()

            is_bold = False
            if para.text.strip():
                non_empty_runs = [run for run in para.runs if run.text.strip()]
                if non_empty_runs:
                    is_bold = all(run.bold for run in non_empty_runs)

            # Check if we're entering Part I context
            if is_in_part_one_context(para.text):
                in_part_one = True
                logger.info(f"Entering Part I context with text: '{para.text}'")

            # Check if text matches any of our special heading patterns
            matches_special_pattern = matches_heading_patterns(para.text)

            if style.startswith("heading"):
                level = style.replace("heading ", "")
                heading_el = etree.SubElement(root, "heading", level=level)
                heading_el.text = para.text
                logger.info(f"Added heading: {para.text}")
            elif matches_special_pattern:
                # If pattern matches, always make it bold regardless of context
                bold_el = etree.SubElement(root, "bold")
                bold_el.text = para.text
                if in_part_one:
                    logger.info(f"Added bold (pattern matched in Part I): {para.text}")
                else:
                    logger.info(f"Added bold (pattern matched): {para.text}")
            elif is_bold:
                bold_el = etree.SubElement(root, "bold")
                bold_el.text = para.text
                logger.info(f"Added bold: {para.text}")
            else:
                para_el = etree.SubElement(root, "paragraph")
                para_el.text = para.text
                logger.info(f"Added paragraph: {para.text[:30]}...")

        elif block.tag.endswith('tbl'):
            table = next((t for t in doc.tables if t._element == block), None)
            if table is None:
                logger.warning("Table block not found in doc.tables.")
                continue
            logger.info("Processing table block.")
            process_table(root, table)


    # Write XML
    try:
        tree = etree.ElementTree(root)
        tree.write(xml_output_path, pretty_print=True, xml_declaration=True, encoding="utf-8")
        logger.info(f"Successfully wrote XML to {xml_output_path}")
    except Exception as e:
        logger.error(f"Failed to write XML to {xml_output_path}: {e}")

def process_all_docx_in_folder(input_folder, output_folder):
    logger.info(f"Processing all DOCX files in folder: {input_folder}")
    """
    Converts all DOCX files in a given folder to XML.
    """

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
        logger.info(f"Created output folder: {output_folder}")


    for filename in os.listdir(input_folder):
        if filename.endswith(".docx"):
            docx_path = os.path.join(input_folder, filename)
            xml_filename = os.path.splitext(filename)[0] + ".xml"
            xml_output_path = os.path.join(output_folder, xml_filename)
            logger.info(f"Processing {docx_path} -> {xml_output_path}")
            print(f"Processing {docx_path} -> {xml_output_path}")
            docx_to_custom_xml(docx_path, xml_output_path)
            print("Done.")
            logger.info(f"Finished processing {docx_path}")

# Example usage:
# Replace with your actual input and output folder paths
#input_directory = "/home/Comptroller_and_Auditor_General/TN_35"
#output_directory = "/home/Comptroller_and_Auditor_General/TN_35_xml"
#process_all_docx_in_folder(input_directory, output_directory)
